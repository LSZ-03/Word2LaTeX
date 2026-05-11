"""DocxParseAgent — orchestrator for word document parsing.

Coordinates:
  1. ooxml_extractor   → sections, paragraphs, formatting
  2. image_extractor   → embedded images
  3. table_extractor   → tables
  4. equation_extractor → OMML equations

Output: a fully populated Manuscript AST node.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional
from uuid import uuid4

from manuscript_compiler.ast.manuscript import Manuscript, ManuscriptMetadata, ManuscriptWarning
from manuscript_compiler.ast.section import SemanticRole
from manuscript_compiler.ast.citation import Citation
from manuscript_compiler.agents.docx_parser.ooxml_extractor import (
    extract_sections,
    extract_title,
    extract_abstract,
)
from manuscript_compiler.agents.docx_parser.image_extractor import extract_images
from manuscript_compiler.agents.docx_parser.table_extractor import extract_tables
from manuscript_compiler.agents.docx_parser.equation_extractor import extract_equations


def run(
    docx_path: str | Path,
    output_dir: Optional[str | Path] = None,
    journal_target: Optional[str] = None,
) -> Manuscript:
    """Run the full docx parsing pipeline.

    Args:
        docx_path: Path to the .docx file.
        output_dir: Directory for extracted assets (images, etc.).
                    Defaults to a 'media/' directory next to the docx.
        journal_target: Optional target journal name.

    Returns:
        A populated Manuscript AST.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    if output_dir is None:
        output_dir = docx_path.parent / f"{docx_path.stem}_media"
    output_dir = Path(output_dir)

    # ── Extract ──────────────────────────────────────────────────────────
    warnings: list[ManuscriptWarning] = []

    # 1. Sections & paragraphs (now returns sections + bibliography)
    sections, bibliography_paragraphs = extract_sections(docx_path)

    # 2. Title & abstract
    from docx import Document as DocxDocument
    doc = DocxDocument(str(docx_path))
    title = extract_title(doc)
    abstract = extract_abstract(doc)

    # 3. Images
    try:
        figures = extract_images(docx_path, output_dir)
    except Exception as e:
        figures = []
        warnings.append(ManuscriptWarning(
            source="docx_parser.image_extractor",
            message=f"Image extraction failed: {e}",
            severity="warning",
        ))

    # 4. Tables
    try:
        tables = extract_tables(docx_path)
    except Exception as e:
        tables = []
        warnings.append(ManuscriptWarning(
            source="docx_parser.table_extractor",
            message=f"Table extraction failed: {e}",
            severity="warning",
        ))

    # 5. Equations — WMF/MathType formulas (returns placeholders)
    try:
        equations, para_texts = extract_equations(docx_path, output_dir=output_dir.parent)
    except Exception as e:
        equations = []
        para_texts = {}
        warnings.append(ManuscriptWarning(
            source="docx_parser.equation_extractor",
            message=f"Equation extraction failed: {e}",
            severity="warning",
        ))

    # 6. Bibliography — convert reference paragraphs to Citation objects
    bibliography = _paragraphs_to_citations(bibliography_paragraphs)

    # 7. Write para_texts (with EQ placeholders) back to paragraph.plain_text
    _write_para_texts(sections, para_texts)

    # 8. Map equations/figures/tables to their paragraph positions
    _map_inline_elements(sections, equations, figures, tables)

    # ── Build Manuscript ─────────────────────────────────────────────────
    metadata = ManuscriptMetadata(
        title=title,
        abstract=abstract,
        journal_target=journal_target,
    )

    manuscript = Manuscript(
        metadata=metadata,
        sections=sections,
        figures=figures,
        tables=tables,
        equations=equations,
        bibliography=bibliography,
        source_file=str(docx_path.resolve()),
        run_id=str(uuid4())[:8],
        pipeline_stage="parsed",
        overall_confidence=1.0,
        warnings=warnings,
    )

    # Try to detect semantic roles for sections (basic heuristic)
    _detect_semantic_roles(manuscript)

    return manuscript


def _write_para_texts(sections, para_texts) -> None:
    """Append EQ placeholders to paragraph plain_text without overwriting content.

    Safe approach: keeps the original python-docx text intact, runs intact,
    only appends \\x00EQ\\x00 markers for the renderer to expand.
    The exact position within the paragraph is not yet precise (TODO),
    but ALL original content, both plain_text and runs, is preserved.
    """
    def _walk(section):
        for para in section.paragraphs:
            if para.paragraph_index is not None and para.paragraph_index in para_texts:
                # Count how many placeholders are needed
                n_placeholders = para_texts[para.paragraph_index].count("\x00EQ\x00")
                if n_placeholders > 0:
                    # Append placeholders to existing text, preserving all original content
                    suffix = " " + " ".join(["\x00EQ\x00"] * n_placeholders)
                    if para.plain_text:
                        para.plain_text = para.plain_text.rstrip() + suffix
                    # DO NOT clear runs — constraint layer needs them for cite/ref processing
        for sub in section.subsections:
            _walk(sub)
    for sec in sections:
        _walk(sec)


def _map_inline_elements(sections, equations, figures, tables) -> None:
    """Map equations to their paragraph positions in the section tree.

    Each equation has a paragraph_index pointing to the docx paragraph it
    belongs to. We iterate the section tree and mark each paragraph with
    which inline elements should appear after it.
    """
    # Build lookup: paragraph_index → equation IDs
    eq_map: dict[int, list[str]] = {}
    for eq in equations:
        if eq.paragraph_index is not None:
            eq_map.setdefault(eq.paragraph_index, []).append(eq.id)
    
    def _walk(section):
        for para in section.paragraphs:
            if para.paragraph_index is not None:
                para.inline_equation_ids = eq_map.get(para.paragraph_index, [])
        for sub in section.subsections:
            _walk(sub)

    for sec in sections:
        _walk(sec)


def _paragraphs_to_citations(paragraphs) -> List[Citation]:
    """Convert reference paragraphs to Citation AST nodes.

    For MVP: store the raw text as-is. A dedicated CitationAgent
    will parse them into structured fields later.
    """
    citations = []
    for i, p in enumerate(paragraphs):
        text = p.plain_text.strip()
        if not text:
            continue
        # Extract first number as key
        key = f"ref_{i + 1}"
        citations.append(Citation(
            key=key,
            title=text[:100],
            entry_type="misc",
            confidence=0.3,  # Low — needs dedicated citation parsing
        ))
    return citations


def _detect_semantic_roles(manuscript: Manuscript) -> None:
    """Basic keyword-based semantic role detection for sections."""
    keyword_map = {
        SemanticRole.INTRODUCTION: {"introduction", "intro", "background"},
        SemanticRole.RELATED_WORK: {"related work", "literature", "prior work"},
        SemanticRole.METHOD: {"method", "approach", "proposed", "architecture", "model"},
        SemanticRole.EXPERIMENTS: {"experiment", "evaluation", "experimental setup"},
        SemanticRole.RESULTS: {"result", "finding", "observation", "analysis"},
        SemanticRole.DISCUSSION: {"discussion", "discuss"},
        SemanticRole.CONCLUSION: {"conclusion", "summary", "future work"},
    }

    def _apply(section):
        title_lower = section.title.lower()
        for role, keywords in keyword_map.items():
            if any(kw in title_lower for kw in keywords):
                section.semantic_role = role
                break
        for sub in section.subsections:
            _apply(sub)

    for sec in manuscript.sections:
        _apply(sec)
