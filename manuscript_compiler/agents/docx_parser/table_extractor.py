"""Table extraction from .docx files.

Uses python-docx's built-in table API to extract table structure
and convert to AST Table nodes.
"""

from __future__ import annotations

from pathlib import Path
import re
from typing import List

from docx import Document as DocxDocument

from manuscript_compiler.ast.table import Table, TableRow, TableCell


def extract_tables(docx_path: str | Path) -> List[Table]:
    """Extract all tables from a .docx file.

    Returns:
        List of Table AST nodes.
    """
    doc = DocxDocument(str(docx_path))
    tables: List[Table] = []

    for i, doc_table in enumerate(doc.tables):
        rows: List[TableRow] = []
        num_cols = 0

        # Find caption for this table
        table_caption = _find_table_caption(docx_path, i)
        # Find notes/abbreviations after this table
        table_notes = _find_table_notes(docx_path, i)

        for row_idx, doc_row in enumerate(doc_table.rows):
            cells: List[TableCell] = []
            for cell in doc_row.cells:
                text = cell.text.strip()
                # Check if this is a header cell (first row heuristic)
                is_header = row_idx == 0

                # Access underlying XML for colspan/rowspan info
                tc = cell._tc
                grid_span = tc.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan")
                v_merge = tc.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge")

                colspan = 1
                if grid_span is not None and grid_span.get("w:val"):
                    try:
                        colspan = int(grid_span.get("w:val"))
                    except (ValueError, TypeError):
                        pass

                rowspan = 1
                if v_merge is not None:
                    val = v_merge.get("w:val")
                    if val and val == "restart":
                        rowspan = 2  # simplified; real merge requires scanning

                tc_cell = TableCell(
                    text=text,
                    colspan=colspan,
                    rowspan=rowspan,
                    is_header=is_header,
                )
                cells.append(tc_cell)

            row = TableRow(cells=cells)
            rows.append(row)
            num_cols = max(num_cols, len(cells))

        table_id = f"tbl_{i + 1}"
        tables.append(Table(
            id=table_id,
            label=f"Table {i + 1}",
            caption=table_caption,
            notes=table_notes,
            rows=rows,
            num_cols=num_cols,
            num_rows=len(rows),
            has_header_row=len(rows) > 0,
        ))

    return tables


def _find_table_notes(docx_path: Path, table_index: int) -> List[str]:
    """Find note/abbreviation paragraphs that follow a table.

    Scans the body after the table element for paragraphs starting
    with Note:/Notes:/Abbreviation:/Abbreviations:, or (*)-style
    statistical note lines.

    Returns:
        List of note text strings.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument(str(docx_path))
    body = doc.element.body
    table_count = -1
    notes: List[str] = []

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "tbl":
            table_count += 1
            if table_count == table_index:
                # Scan following paragraphs for notes
                sibling = child.getnext()
                paragraphs_scanned = 0
                while sibling is not None and paragraphs_scanned < 3:
                    sibling_tag = sibling.tag.split("}")[-1] if "}" in sibling.tag else sibling.tag
                    if sibling_tag == "p":
                        texts = []
                        for t in sibling.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                            if t.text:
                                texts.append(t.text)
                        text = "".join(texts).strip()

                        if re.match(r"^(Note|Notes|Abbreviation|Abbreviations|\*)", text, re.IGNORECASE):
                            notes.append(text)
                        elif notes:
                            # Continuation of previous note
                            notes[-1] += " " + text
                        else:
                            # Not a note, stop scanning
                            break
                    elif sibling_tag == "tbl":
                        break  # Next table starts
                    paragraphs_scanned += 1
                    sibling = sibling.getnext()
                break
    return notes


def _find_table_caption(docx_path: Path, table_index: int) -> str:
    """Try to find a caption paragraph preceding the table.

    Common pattern: "Table 1: ..." appears as a paragraph right before
    the table in the document body.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument(str(docx_path))

    # Walk through document body elements in order
    body = doc.element.body
    table_count = -1
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "tbl":
            table_count += 1
            if table_count == table_index:
                # Look at previous sibling for caption
                prev = child.getprevious()
                if prev is not None:
                    prev_tag = prev.tag.split("}")[-1] if "}" in prev.tag else prev.tag
                    if prev_tag == "p":
                        texts = []
                        for t in prev.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                            if t.text:
                                texts.append(t.text)
                        caption = "".join(texts).strip()
                        if caption and ("Table" in caption or "TABLE" in caption):
                            return caption
    return ""
