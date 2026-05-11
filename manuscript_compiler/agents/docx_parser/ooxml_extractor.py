"""OOXML paragraph & section extraction from .docx files.

Uses python-docx to parse paragraphs, headings, inline formatting,
and builds a section tree from heading levels.

Handles:
  - Word Heading styles (Heading 1, Heading 2, ...)
  - Heuristic text-based headings ("1. Introduction", "Abstract:", etc.)
  - Reference/bibliography sections (numbered refs are NOT headings)
  - List items that look like headings but aren't ("1) ...", "(1) ...")
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph as DocxParagraph

from manuscript_compiler.ast.section import (
    Section,
    Paragraph,
    TextRun,
    SemanticRole,
)


# ── Heading detection patterns ─────────────────────────────────────────

# Positive: text patterns that ARE headings
_HEADING_PATTERNS = [
    # (word) style headings: "1. Introduction", "3.1 Method", "2.2.1 Details"
    re.compile(r"^(\d+(?:\.\d+)*)\.?\s+[A-Z][a-zA-Z\s\-/]+$"),
    # Named section markers (case-insensitive at start of line)
    re.compile(r"^(Abstract|Keywords|Introduction|Related\s+Work|Method(ology)?|"
               r"Experiments?\s*(and|&)?\s*(Analysis|Results?)?|"
               r"Results?\s*(and|&)?\s*(Analysis|Discussion)?|"
               r"Discussion|Conclusion|Appendix|References|Acknowledgments?|"
               r"Supplementary\s+Material)(\s|:|\b)",
               re.IGNORECASE),
    # Roman numerals: "I. INTRODUCTION", "II. BACKGROUND"
    re.compile(r"^(I|II|III|IV|V|VI|VII|VIII|IX|X)\.\s+[A-Z]"),
]

# Negative: text patterns that look like headings but AREN'T
_HEADING_NEGATIVE_PATTERNS = [
    # Reference entries: "1. Author, Firstname. Title..."
    re.compile(r"^\d+\.\s+[A-Z][a-z]+,\s+[A-Z][a-z]+"),
    # Figure/Table captions: "Fig. 1.", "Figure 2:", "TABLE I"
    re.compile(r"^(Fig(ure)?|Table|TABLE)\s+\d+"),
    # Numbered list items: "1) ...", "(1) ...", "(a) ...", "a) ..."
    re.compile(r"^\d+\)\s+"),
    re.compile(r"^\(\d+\)\s+"),
    re.compile(r"^[a-z]\)\s+"),
    re.compile(r"^\([a-z]\)\s+"),
    # Bullet-like dashes at start
    re.compile(r"^[-•·]\s+"),
    # DOI / URL lines
    re.compile(r"^https?://|^doi:"),
]


def _detect_heading_level(text: str, style_name: str | None) -> int:
    """Detect heading level from style name or text pattern.

    Only checks the FIRST LINE of the text (pre-\\n) to handle
    documents where heading and body are in the same paragraph.

    Returns 0 if not a heading, 1-9 for heading level.
    """
    # Only look at the first line — body text after \\n is not a heading
    first_line = text.strip().split("\n")[0].strip()
    if not first_line:
        return 0

    # ── Method 1: Word Heading styles ──
    if style_name:
        sn = style_name.lower()
        if sn.startswith("heading"):
            level_str = sn.replace("heading", "").replace(" ", "").strip()
            try:
                return int(level_str) if level_str else 1
            except ValueError:
                return 1
        if "title" in sn:
            return 1

    # ── Method 2: Heuristic text patterns (on first line only) ──

    # First check negative patterns (fast rejection)
    for neg in _HEADING_NEGATIVE_PATTERNS:
        if neg.match(first_line):
            return 0

    # Then check positive patterns
    for pattern in _HEADING_PATTERNS:
        m = pattern.match(first_line)
        if m:
            num_part = m.group(1)
            if num_part and num_part[0].isdigit():
                dots = num_part.count(".")
                return min(dots + 1, 9)
            return 1

    # ── Method 3: Single short line with all caps ──
    if len(first_line.split()) <= 5 and first_line.isupper():
        return 1

    return 0


def _extract_heading_and_body(text: str, heading_level: int) -> tuple[str, str]:
    """Split a heading paragraph into heading title and remaining body text.

    Handles cases like:
      "Abstract: The field of remote sensing..." → ("Abstract", "The field of...")
      "Keywords: A; B; C" → ("Keywords", "A; B; C")
      "1. Introduction" → ("1. Introduction", "")
      "2.3. Sample Selection\n\tIn detection..." → ("2.3. Sample Selection", "In detection...")

    Returns:
        (heading_title, body_text)
    """
    first_line = text.split("\n")[0].strip()

    # Case 1: Named heading with colon — "Abstract: ..." or "Keywords: ..."
    # Detect by checking if heading_level comes from a named (non-numeric) match
    if heading_level > 0:
        for pattern in _HEADING_PATTERNS:
            m = pattern.match(first_line)
            if m:
                num_part = m.group(1)
                if num_part and num_part[0].isdigit():
                    # Numeric heading like "1. Introduction" — NO colon splitting
                    heading_title = first_line
                    body_after_heading = first_line[len(m.group(0)):].strip()
                    # body_after_heading is empty for well-formed numeric headings
                    break
                else:
                    # Named heading like "Abstract:" — check for colon
                    keyword = m.group(1).strip()
                    if ":" in first_line and first_line.index(":") < len(keyword) + 3:
                        # "Abstract: the rest..." → heading = "Abstract", body = "the rest..."
                        colon_idx = first_line.index(":")
                        heading_title = first_line[:colon_idx].strip()
                        body_after_heading = first_line[colon_idx + 1:].strip()
                    else:
                        heading_title = keyword
                        body_after_heading = first_line[len(keyword):].strip().lstrip(": ")
                    break
        else:
            # All-caps heading or other — use full first line
            heading_title = first_line
            body_after_heading = ""
    else:
        heading_title = first_line
        body_after_heading = ""

    # Combine with lines after \\n
    rest_lines = [l.strip() for l in text.split("\n")[1:] if l.strip()]
    body_parts = []
    if body_after_heading:
        body_parts.append(body_after_heading)
    body_parts.extend(rest_lines)

    return heading_title, "\n".join(body_parts)


def _is_reference_entry(text: str) -> bool:
    """Check if a paragraph looks like a bibliography reference entry."""
    if not text:
        return False
    t = text.strip()
    # Pattern: "N. Surname, F. Title..." or "[N] Surname..."
    ref_patterns = [
        re.compile(r"^\d+\.\s+[A-Z][a-z]+,\s+[A-Z]"),       # "1. Cheng, G."
        re.compile(r"^\[\d+\]\s+[A-Z][a-z]+"),                # "[1] Author..."
        re.compile(r"^\d+\.\s+[‘’\"\"][A-Z]"),                  # "1. 'Title..." 
    ]
    return any(p.match(t) for p in ref_patterns)


# ── Main extraction API ────────────────────────────────────────────────


def extract_sections(docx_path: str | Path) -> tuple:
    """Parse a .docx file and build a section tree + bibliography.

    Returns:
        (sections: List[Section], bibliography: List[Paragraph])
    """
    doc = DocxDocument(str(docx_path))
    paragraphs: List[Paragraph] = []
    for i, para in enumerate(doc.paragraphs):
        paragraphs.extend(_parse_paragraph(para, i))

    return _build_section_tree(paragraphs)


def extract_abstract(doc: DocxDocument) -> str:
    """Extract the abstract text from the document.

    Strategy: find a paragraph that starts with "Abstract" (case-insensitive),
    then collect its body text (split by colon/换行). Stop at the next heading
    or a Keywords/Introduction line.
    """
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""

        # Detect abstract start
        is_abstract_style = "abstract" in style
        is_abstract_text = text.lower().startswith("abstract")

        if is_abstract_style or is_abstract_text:
            # If the paragraph contains "Abstract: body...", extract body
            lower = text.lower()
            if "abstract" in lower:
                idx = lower.index("abstract") + len("abstract")
                body = text[idx:].lstrip(": .\t\n").strip()
                # Stop at "Keywords", "1. Introduction", etc.
                # Take only the first contentful block
                parts = []
                for line in body.split("\n"):
                    line = line.strip()
                    if not line:
                        break
                    if line.lower().startswith("keywords"):
                        break
                    if re.match(r"^\d+\.\s", line):
                        break
                    parts.append(line)
                return "\n".join(parts).strip()
            return text

    return ""


def extract_title(doc: DocxDocument) -> str:
    """Extract manuscript title."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""
        if "title" in style:
            return text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            return text
    return ""


# ── Paragraph parsing ──────────────────────────────────────────────────


def _parse_paragraph(para: DocxParagraph, index: int) -> List[Paragraph]:
    """Convert a python-docx paragraph to one or more AST Paragraphs.

    Returns a LIST because one Word paragraph may contain both a heading
    and body text (e.g. "Abstract: body text..." or "2.3. Title\n\tbody...").

    Returns empty list if the paragraph has no meaningful text.
    """
    text = para.text.strip()
    if not text:
        return []

    style_name = para.style.name if para.style else None
    heading_level = _detect_heading_level(text, style_name)

    if heading_level > 0:
        heading_title, body_text = _extract_heading_and_body(text, heading_level)
        heading_p = Paragraph(
            runs=[TextRun(text=heading_title)],
            plain_text=heading_title,
            style_name=style_name,
            is_heading=True,
            heading_level=heading_level,
            paragraph_index=index,
        )
        if body_text:
            body_p = Paragraph(
                plain_text=body_text,
                style_name=style_name,
                paragraph_index=None,
            )
            return [heading_p, body_p]
        return [heading_p]

    # ── Normal: non-heading paragraph ──
    runs = []
    for run in para.runs:
        if not run.text.strip():
            continue
        tr = TextRun(
            text=run.text,
            bold=run.bold or False,
            italic=run.italic or False,
            underline=run.underline or False,
            font_size=run.font.size.pt if run.font.size else None,
            font_name=run.font.name,
        )
        runs.append(tr)

    align_map = {None: None, 0: "left", 1: "center", 2: "right", 3: "both"}
    alignment = align_map.get(para.alignment)

    return [Paragraph(
        runs=runs,
        plain_text=text,
        style_name=style_name,
        alignment=alignment,
        is_heading=False,
        heading_level=0,
        paragraph_index=index,
    )]


# ── Section tree builder ───────────────────────────────────────────────


def _build_section_tree(paragraphs: List[Paragraph]) -> tuple:
    """Build nested section tree from flat paragraph list.

    Detects bibliography sections and extracts reference entries
    as separate bibliography paragraphs rather than subsections.

    Returns:
        (root_sections: List[Section], bibliography: List[Paragraph])
    """
    root_sections: List[Section] = []
    stack: List[Section] = []
    section_counter = 0
    bibliography: List[Paragraph] = []
    in_bibliography = False

    def _make_section(title: str, level: int) -> Section:
        nonlocal section_counter
        section_counter += 1
        return Section(
            id=f"sec_{section_counter}",
            title=title,
            level=level,
        )

    for para in paragraphs:
        # ── Detect bibliography mode ──
        if para.is_heading and para.text.strip().lower().startswith("references"):
            # Create the References section heading
            sec = _make_section(para.text, para.heading_level)
            while stack and stack[-1].level >= para.heading_level:
                stack.pop()
            if stack:
                stack[-1].subsections.append(sec)
            else:
                root_sections.append(sec)
            stack.append(sec)
            in_bibliography = True
            continue

        # In bibliography mode: collect reference entries
        if in_bibliography:
            if para.is_heading:
                # Another real heading found — exit bibliography mode
                in_bibliography = False
                # Process normally below
            else:
                bibliography.append(para)
                continue

        # ── Normal section building ──
        if para.is_heading:
            sec = _make_section(para.text, para.heading_level)

            while stack and stack[-1].level >= para.heading_level:
                stack.pop()

            if stack:
                stack[-1].subsections.append(sec)
            else:
                root_sections.append(sec)
            stack.append(sec)
        else:
            if stack:
                stack[-1].paragraphs.append(para)
            else:
                if not root_sections or root_sections[0].title != "__preamble__":
                    preamble = _make_section("__preamble__", 0)
                    root_sections.insert(0, preamble)
                    stack.append(preamble)
                stack[-1].paragraphs.append(para)

    return root_sections, bibliography
